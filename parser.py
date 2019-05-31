from glob import glob
from subprocess import check_output, CalledProcessError
from typing import Union, Dict, List
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
import numpy as np
import pandas as pd
from bs4 import BeautifulSoup


class PathTools:

    def filepaths(self, folder: Union[Path, str]) -> list:
        ''' folder is a text that follows glob syntax '''
        return [self.filepath(f) for f in glob(folder)]

    def filepath(self, filepath: Union[Path, str]) -> str:
        path = Path(filepath)
        if path.is_file():
            path = str(path.absolute())
            return path
        print(f'WARNING ::: {filepath} is not a file')


class TextTools:

    def remove_unicode(self, string: str) -> str:
        '''
        Removes unicode by converting the string into native bytes and then into
        ascii to force unicode out because it can't be converted.

        Args: string (str): string containing unicode
        Returns: string without unicode
        '''
        # we dont want to remove space unicode
        string = string.replace('\xa0', ' ')
        # convert to bytes to get the unicde to show up properly
        byte_string = bytes(string, 'utf-8')
        # removes unicode and ignores errors from the unicode
        byte_str_without_unicode = byte_string.decode("ascii", "ignore")

    def count_head_spaces(self, string):
        return len(string) - len(string.lstrip(' '))

    def count_tail_spaces(self, string):
        return len(string) - len(string.rstrip(' '))

    def create_grid(self, col_size, indx_size, init_value=None):
        col_size += 1
        indx_size += 1
        return np.array([[init_value] * col_size] * indx_size)


class Docx(PathTools, TextTools):

    def __init__(self, document_path):
        self.docx = self.open_docx(document_path)

    def open_docx(self, filepath: Union[Path, str]) -> Document:
        ''' Open Docx Document '''
        with open(self.filepath(filepath), 'rb') as infile:
            return Document(infile)

    def get_header_text(self) -> list:
        '''
        BUGS: if header is multiple lines or if you signed your name on the same level
            as the header it might not get it all or attach too much.
        '''

        def find_header_text(paragraphs: List[Paragraph]) -> str:
            '''
            Args: paragraphs(list): segment of the document
            Returns: str if header is found else None
            '''
            # print(paragraphs)
            header_text = []
            for header in paragraphs:
                if not header.text:
                    continue
                elif header.text.isdigit():
                    continue
                else:
                    header_text.append(header.text)

        # If file has proper header it will be in the header segment of document
        header_text = [find_header_text(sec.header.paragraphs) for sec in self.docx.sections]
        header_text = [t for t in header_text if t]
        return header_text

    def get_body_title(self) -> str:
        '''
        BUGS: if header is multiple lines or if you signed your name on the same level
            as the header it might not get it all or attach too much.
        '''

        def find_first_text(paragraphs: List[Paragraph]) -> str:
            '''
            Args: paragraphs(list): segment of the document
            Returns: str if header is found else None
            '''
            for header in paragraphs:
                if not header.text:
                    continue
                elif header.text.isdigit():
                    continue
                else:
                    return header.text

        first_text = find_first_text(self.docx.paragraphs)
        return first_text

class Pandoc(PathTools, TextTools):

    def __init__(self, infile=None):
        self.infile = infile
        self.load(infile) if infile else None

    def shell(self, string: str) -> str:
        '''
        check_output is the only one to behave like a real shell

        Args: string (str): shell script
        Returns: whatever the shell script was intended to return; else 0
        '''
        try:
            return check_output(string, shell=True).decode('UTF-8')
        except CalledProcessError:
            print(f'script >{string}< returned an error')

    def load(self, filepath):
        filepath = self.filepath(filepath)
        filepath = filepath.replace(' ', '\ ')
        content = self.shell(f'pandoc -s {filepath}')
        self.content = content
        self.soup = BeautifulSoup(content, 'lxml')

    def get_body_title(self):
        for line in self.soup.body.text.splitlines():
            if line.strip():
                return line


def example():

    for filepath in PathTools().filepaths('./SCG/data/**'):
        docx = Docx(filepath)
        print(f'python-docx ::: {docx.get_body_title()}')
        doc = Pandoc(filepath)
        print(f'pandoc      ::: {doc.get_body_title()}')


if __name__ == '__main__':
    example()
